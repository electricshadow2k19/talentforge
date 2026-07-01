"""Section-based DOCX builder — mirrors frontend resumeParser logic."""

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BRAND = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x50, 0x50, 0x50)

SECTION_HEADERS = [
    "PROFESSIONAL SUMMARY", "SUMMARY", "PROFILE", "OBJECTIVE",
    "TECHNICAL SKILLS", "SKILLS", "CORE COMPETENCIES",
    "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EXPERIENCE",
    "EDUCATION", "CERTIFICATIONS", "PROJECTS",
]

CONTACT_RE = re.compile(r"@|linkedin|github|phone:|email:|tel:|http|www\.|^\+\d", re.I)
TITLE_RE = re.compile(
    r"engineer|developer|architect|analyst|manager|consultant|administrator|specialist|\blead\b",
    re.I,
)


def _normalize_header(line: str) -> str | None:
    t = line.strip()
    if not t:
        return None
    bare = t.rstrip(":")
    for h in SECTION_HEADERS:
        if re.match(rf"^{re.escape(h)}\s*:?\s*$", bare, re.I):
            return h
    if re.match(r"^TECHNICAL\s+SKILLS?\s*:", t, re.I):
        return "TECHNICAL SKILLS"
    if re.match(r"^SKILLS?\s*:", t, re.I) and len(t) < 40:
        return "TECHNICAL SKILLS"
    return None


def _is_job_header(line: str) -> bool:
    t = line.strip()
    if len(t) > 130 or t.startswith(("-", "•")):
        return False
    if re.search(r"\d{4}\s*[-–—]\s*(?:\d{4}|present)", t, re.I):
        return True
    if "|" in t and len(t) < 100 and not t.endswith("."):
        parts = [p.strip() for p in t.split("|")]
        return len(parts) >= 2 and len(parts[0]) < 50
    return False


def _is_bullet(line: str) -> bool:
    return bool(re.match(r"^[\-\*•]\s+", line.strip()))


def _parse_resume(text: str) -> dict:
    lines = text.replace("\r\n", "\n").split("\n")
    sections: list[dict] = []
    current: dict | None = None
    header_lines: list[str] = []

    for raw in lines:
        line = raw.rstrip()
        header = _normalize_header(line)
        if header:
            current = {"title": header, "lines": []}
            sections.append(current)
            continue
        if re.match(r"^(TECHNICAL\s+)?SKILLS?\s*[:|]", line.strip(), re.I):
            body = re.sub(r"^(TECHNICAL\s+)?SKILLS?\s*[:|]\s*", "", line.strip(), flags=re.I)
            if not current or current["title"] != "TECHNICAL SKILLS":
                current = {"title": "TECHNICAL SKILLS", "lines": []}
                sections.append(current)
            if body:
                current["lines"].append(body)
            continue
        if current:
            if line.strip():
                current["lines"].append(line.strip())
        elif line.strip():
            header_lines.append(line.strip())

    name, title, contact, pre = "Candidate", "", [], []
    for i, line in enumerate(header_lines):
        if i == 0 and len(line) < 70 and not CONTACT_RE.search(line):
            name = line
            continue
        if not title and len(line) < 90 and TITLE_RE.search(line) and not CONTACT_RE.search(line):
            title = line
            continue
        if CONTACT_RE.search(line) or (contact and len(line) < 120):
            contact.append(line)
            continue
        pre.append(line)

    if pre and sections and re.search(r"SUMMARY|PROFILE|OBJECTIVE", sections[0]["title"], re.I):
        sections[0]["lines"] = pre + sections[0]["lines"]
    elif pre:
        sections.insert(0, {"title": "PROFESSIONAL SUMMARY", "lines": pre})

    return {"name": name, "title": title, "contact": contact, "sections": sections}


def _section_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "2E75B6")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def resume_text_to_docx_bytes(text: str) -> bytes:
    parsed = _parse_resume(text)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(parsed["name"])
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = BRAND

    if parsed["title"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(parsed["title"])
        r.italic = True
        r.font.size = Pt(12)
        r.font.color.rgb = GRAY

    if parsed["contact"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("  |  ".join(parsed["contact"]))
        r.font.size = Pt(10)
        r.font.color.rgb = GRAY

    doc.add_paragraph("")

    for section in parsed["sections"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        _section_border(p)
        r = p.add_run(section["title"].upper())
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = BRAND

        is_skills = bool(re.search(r"SKILLS|COMPETENCIES", section["title"], re.I))
        for line in section["lines"]:
            if _is_bullet(line):
                doc.add_paragraph(re.sub(r"^[\-\*•]\s+", "", line.strip()), style="List Bullet")
            elif _is_job_header(line):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                r = p.add_run(line)
                r.bold = True
                r.font.color.rgb = BRAND
            elif is_skills and "|" in line:
                skills = [s.strip() for s in re.split(r"\||,", line) if s.strip()]
                doc.add_paragraph("    ".join(f"• {s}" for s in skills))
            else:
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
